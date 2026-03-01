# bot/ocr.py

import os
import logging
import platform
import tempfile
from typing import Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor

import fitz  # PyMuPDF

# Подавляем лишние логи от EasyOCR
logging.getLogger('easyocr').setLevel(logging.WARNING)

# EasyOCR для распознавания текста с изображений
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    # Создаем глобальный reader для русского и английского
    easyocr_reader = easyocr.Reader(['ru', 'en'], gpu=False, model_storage_directory='temp/easyocr', verbose=False)
    logging.info("✅ EasyOCR инициализирован")
except ImportError:
    EASYOCR_AVAILABLE = False
    logging.warning("⚠️ EasyOCR не установлен. Установите: pip install easyocr")

# Для поддержки Word документов
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("⚠️ python-docx не установлен. Word документы не поддерживаются.")

logger = logging.getLogger(__name__)

# Пул потоков для CPU-интенсивных операций
executor = ThreadPoolExecutor(max_workers=2)

async def extract_text_from_pdf(pdf_path: str) -> str:
    """Извлечение текста из PDF"""
    text_parts = []
    
    try:
        pdf_document = fitz.open(pdf_path)
        logger.info(f"PDF открыт, страниц: {len(pdf_document)}")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Пробуем извлечь встроенный текст
            page_text = page.get_text()
            
            if page_text.strip():
                logger.info(f"Страница {page_num + 1}: найден текст ({len(page_text)} символов)")
                text_parts.append(f"--- Страница {page_num + 1} ---\n{page_text}")
            else:
                # Если текста нет, делаем пометку
                text_parts.append(f"--- Страница {page_num + 1} ---\n[Текст отсутствует или не распознан]")
        
        pdf_document.close()
        
    except Exception as e:
        logger.error(f"Ошибка PDF: {e}")
        return f"Ошибка обработки PDF: {str(e)}"
    
    return "\n\n".join(text_parts)

async def extract_text_from_image(image_path: str) -> str:
    """Извлечение текста из изображения с помощью EasyOCR"""
    if not EASYOCR_AVAILABLE:
        return "EasyOCR не установлен. Установите: pip install easyocr"
    
    try:
        # Используем EasyOCR для распознавания
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            executor,
            lambda: easyocr_reader.readtext(image_path, detail=0, paragraph=True)
        )
        
        if result:
            text = '\n'.join(result)
            logger.info(f"EasyOCR распознал {len(text)} символов")
            return text.strip()
        else:
            return "На изображении не найден текст"
        
    except Exception as e:
        logger.error(f"Ошибка EasyOCR: {e}")
        return f"Ошибка распознавания: {str(e)}"

async def extract_text_from_docx(docx_path: str) -> str:
    """Извлечение текста из Word документа"""
    if not DOCX_SUPPORT:
        return "Поддержка Word документов не установлена. Установите: pip install python-docx"
    
    try:
        doc = DocxDocument(docx_path)
        text_parts = []
        
        # Извлекаем текст из параграфов
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        result = '\n'.join(text_parts)
        logger.info(f"Word документ: {len(result)} символов")
        return result
        
    except Exception as e:
        logger.error(f"Ошибка обработки Word: {e}")
        return f"Ошибка обработки Word: {str(e)}"

async def extract_text(file_path: str, mime_type: str) -> str:
    """Основная функция извлечения текста"""
    logger.info(f"Извлечение текста из: {file_path}")
    
    if not os.path.exists(file_path):
        return "Ошибка: файл не найден"
    
    try:
        # PDF
        if mime_type == "application/pdf" or file_path.lower().endswith('.pdf'):
            return await extract_text_from_pdf(file_path)
        
        # Word
        elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or \
             file_path.lower().endswith(('.doc', '.docx')):
            return await extract_text_from_docx(file_path)
        
        # Изображения
        elif mime_type.startswith('image/') or file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
            return await extract_text_from_image(file_path)
        
        else:
            return f"Неподдерживаемый тип файла: {mime_type}"
            
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return f"Ошибка: {str(e)}"